"""
DOCX text and metadata extraction.

Wraps ``python-docx`` behind a small, typed interface so the rest of
the application never depends on the underlying parsing library
directly. Extracts paragraph and table text (joined into a single
document body) along with whatever author/title/creation-date metadata
the document's core properties provide.
"""

from __future__ import annotations

import logging
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import EmptyDocumentError, TextExtractionError

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class DocxExtractionResult:
    """
    The result of extracting text and metadata from a DOCX file.

    Attributes:
        text: The full extracted text (paragraphs and table cells, in
            document order), joined by blank lines.
        author: The document author, if present in the core properties.
        title: The document title, if present in the core properties.
        created_at: The document creation timestamp, if present in the core properties.
        extra_metadata: Any additional metadata fields (e.g. subject, keywords).
    """

    text: str
    author: Optional[str] = None
    title: Optional[str] = None
    created_at: Optional[datetime] = None
    extra_metadata: dict[str, str] = field(default_factory=dict)


class DocxLoader:
    """Extracts text and metadata from ``.docx`` files using ``python-docx``."""

    def load(self, file_path: Path) -> DocxExtractionResult:
        """
        Extract text and metadata from a DOCX file.

        Args:
            file_path: Path to the DOCX file on disk.

        Returns:
            A :class:`DocxExtractionResult` with the extracted text and metadata.

        Raises:
            TextExtractionError: If the file does not exist or cannot be
                parsed as a valid DOCX package.
            EmptyDocumentError: If the document contains no extractable text.
        """
        if not file_path.exists():
            raise TextExtractionError(str(file_path), "file does not exist.")

        try:
            document = DocxDocument(str(file_path))
        except (PackageNotFoundError, zipfile.BadZipFile, OSError, ValueError) as exc:
            raise TextExtractionError(str(file_path), f"failed to open DOCX: {exc}") from exc

        text = self._extract_text(document)
        if not text.strip():
            raise EmptyDocumentError(file_path.name)

        author, title, created_at, extra_metadata = self._extract_metadata(document)

        logger.info("Extracted %d character(s) from DOCX '%s'.", len(text), file_path.name)
        return DocxExtractionResult(
            text=text,
            author=author,
            title=title,
            created_at=created_at,
            extra_metadata=extra_metadata,
        )

    @staticmethod
    def _extract_text(document: DocxDocument) -> str:
        """
        Extract all paragraph and table cell text from a DOCX document,
        in document order.

        Args:
            document: The opened DOCX document.

        Returns:
            The extracted text, with each paragraph/cell on its own
            line and blank lines removed.
        """
        segments: list[str] = []

        for paragraph in document.paragraphs:
            stripped = paragraph.text.strip()
            if stripped:
                segments.append(stripped)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    stripped = cell.text.strip()
                    if stripped:
                        segments.append(stripped)

        return "\n\n".join(segments)

    @staticmethod
    def _extract_metadata(
        document: DocxDocument,
    ) -> tuple[Optional[str], Optional[str], Optional[datetime], dict[str, str]]:
        """
        Extract author/title/creation-date/extra metadata from a DOCX
        document's core properties.

        Args:
            document: The opened DOCX document.

        Returns:
            An ``(author, title, created_at, extra_metadata)`` tuple.
        """
        core_properties = document.core_properties

        author = core_properties.author or None
        title = core_properties.title or None
        created_at = core_properties.created

        extra_metadata = {
            key: value
            for key, value in {
                "subject": core_properties.subject,
                "keywords": core_properties.keywords,
                "category": core_properties.category,
                "comments": core_properties.comments,
            }.items()
            if value
        }

        return author, title, created_at, extra_metadata